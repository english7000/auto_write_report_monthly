#coding=utf-8
#！/usr/bin/env python3

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
doc = docx.Document()
font_name = u'宋体'



# '''定义2个字体的格式'''

obj_styles = doc.styles
obj_Titlestyle = obj_styles.add_style('TitleStyle', WD_STYLE_TYPE.CHARACTER)
obj_Titlestyle.font.size = Pt(20)
obj_Titlestyle.font.name = u'宋体'

obj_NormalStyle = obj_styles.add_style("NormalStyle", WD_STYLE_TYPE.CHARACTER)
obj_NormalStyle.font.size = Pt(15)
obj_NormalStyle.font.name = u'宋体'

obj_ParagraphStyle = obj_styles.add_style('PSTYLE', WD_STYLE_TYPE.PARAGRAPH)
obj_ParagraphStyle.font.size = Pt(18)
obj_ParagraphStyle.font.name = u'宋体'

# '''定义2个增加paragraph的方法，分别为这2个方法定义2个不同是的格式'''

paragraph1 = doc.add_paragraph()
paragraph_center = doc.add_paragraph()      #此段落格式居中
paragraph_center.alignment = WD_ALIGN_PARAGRAPH.CENTER


paragraph1.add_run('巡检记录： 已检\n设备型号： Cisco 4506\n设备描述：多层交换机\n',style = 'TitleStyle').bold = True


paragraph_center. add_run('设备基本信息',style='NormalStyle').bold = True


table_1 = doc.add_table(rows=8, cols=4, style='Table Grid')

row0_name_list = ['设备品牌','设备名称','管理地址','软件版本','端口数量','CPU使用率','设备状态灯','设备报错Error Log']
row2_name_list = ['序列号','设备型号','内存大小','放置位置','端口状态','内存利用率','风扇状况','连通性']
# rows = table_21.row                             #获取所有行
# rows[1].cells[0].text = row_name_list[1]        #给第1行第0个单元格赋值


rows = table_1.rows
for i,j,z in zip(rows,row0_name_list,row2_name_list):
    i.cells[0].text = j
    i.cells[2].text = z

table_2 = doc.add_table(rows=4, cols=1)
table_2_cells = []
for i in table_2.rows:          #把table2的单元格存进table_2_cell 列表以便后续调用
    for cell in i.cells:
        table_2_cells.append(cell)

print(dir(table_2_cells[0]))
table_2_cells[0].add_paragraph('故障描述:', style='PSTYLE').bold = True     #保留cell1和cell3为结果填写
# table_2_cells[1].add_paragraph('', style='PSTYLE')
table_2_cells[2].add_paragraph('巡检结果:', style='PSTYLE').bold = True
# table_2_cells[3].add_paragraph('', style='PSTYLE')



doc.save('text2format.docx')